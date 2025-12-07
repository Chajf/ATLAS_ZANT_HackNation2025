from langchain_ollama import ChatOllama
from langchain_groq import ChatGroq
from schemas import (
    InjuryEvaluationResponse, 
    InjuryDescriptionRequest, 
    ComponentEvaluation, 
    PDFExtractionResponse, 
    DOCXExplanationResponse,
    ValidationIssue,
    DataComparisonResponse,
    AccidentNotificationRequest,
    InjuredStatementRequest,
    OfficeAssessmentResponse,
    CriterionAssessment,
    JustificationRequest,
    JustificationResponse,
    AccidentCardRequest
)
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
import json
import os
from pypdf import PdfReader
from docx import Document
import io
import re
from typing import BinaryIO, List
from datetime import datetime
import fitz  # PyMuPDF

# llm = ChatOllama(
#     model="qwen3:latest",
#     num_ctx=12000,
#     reasoning=True
# )

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

llm = ChatGroq(
    model_name="openai/gpt-oss-20b",
    reasoning_effort="low",
    temperature=0.1,
    max_retries=2
)

llm_llama = ChatGroq(
    model_name="llama-3.3-70b-versatile",
    temperature=0,
    max_retries=2
)

with open("injury_evaluation_prompt.txt", "r", encoding="utf-8") as file:
    injury_evaluation_prompt = file.read()

with open("pdf_extraction_prompt.txt", "r", encoding="utf-8") as file:
    pdf_extraction_prompt = file.read()

with open("office_assessment_prompt.txt", "r", encoding="utf-8") as file:
    office_assessment_prompt = file.read()

with open("justification_prompt.txt", "r", encoding="utf-8") as file:
    justification_prompt = file.read()

def evaluate_injury_description(user_msg: str) -> InjuryEvaluationResponse:
    try:
        structured_llm = llm.with_structured_output(
            schema=InjuryEvaluationResponse,
            method="json_mode"
        )

        messages = [
            SystemMessage(content=injury_evaluation_prompt),
            HumanMessage(content=user_msg)
        ]

        ans = structured_llm.invoke(messages)
        return ans
    
    except Exception as e:
        print(f"Error in evaluate_injury_description: {e}")
        # Zwróć domyślną odpowiedź w przypadku błędu
        return InjuryEvaluationResponse(
            When=ComponentEvaluation(
                Status="danger",
                Description="Wystąpił błąd podczas analizy. Spróbuj ponownie lub skontaktuj się z administratorem."
            ),
            Where=ComponentEvaluation(
                Status="danger",
                Description="Wystąpił błąd podczas analizy."
            ),
            Doing=ComponentEvaluation(
                Status="danger",
                Description="Wystąpił błąd podczas analizy."
            ),
            How=ComponentEvaluation(
                Status="danger",
                Description="Wystąpił błąd podczas analizy."
            ),
            Why=ComponentEvaluation(
                Status="danger",
                Description="Wystąpił błąd podczas analizy."
            ),
            Injury=ComponentEvaluation(
                Status="danger",
                Description="Wystąpił błąd podczas analizy."
            )
        )


def assess_workplace_accident(description: str) -> OfficeAssessmentResponse:
    """
    Office worker assessment - evaluate if event qualifies as workplace accident
    according to Polish labor law criteria
    """
    try:
        structured_llm = llm.with_structured_output(
            schema=OfficeAssessmentResponse,
            method="json_mode"
        )

        messages = [
            SystemMessage(content=office_assessment_prompt),
            HumanMessage(content=f"Oceń następujące zdarzenie pod kątem spełnienia kryteriów wypadku przy pracy. Zwróć odpowiedź w formacie JSON zgodnym z wymaganym schematem.\n\nOpis zdarzenia:\n{description}")
        ]

        result = structured_llm.invoke(messages)
        return result
    
    except Exception as e:
        print(f"Error in assess_workplace_accident: {e}")
        # Return default error response
        return OfficeAssessmentResponse(
            sudden=CriterionAssessment(
                status="danger",
                description="Wystąpił błąd podczas analizy. Spróbuj ponownie lub skontaktuj się z administratorem."
            ),
            external=CriterionAssessment(
                status="danger",
                description="Wystąpił błąd podczas analizy systemu."
            ),
            work=CriterionAssessment(
                status="danger",
                description="Wystąpił błąd podczas analizy systemu."
            ),
            injury=CriterionAssessment(
                status="danger",
                description="Wystąpił błąd podczas analizy systemu."
            )
        )


def generate_decision_justification(decision: str, assessment: OfficeAssessmentResponse, validation_issues: List = None) -> str:
    """
    Generate professional justification for office worker's decision
    based on assessment of all four criteria and document validation issues
    """
    try:
        from schemas import JustificationResponse
        
        structured_llm = llm_llama.with_structured_output(
            schema=JustificationResponse,
            method="json_mode"
        )
        
        # Prepare assessment summary for the prompt
        status_map = {
            "ok": "spełnione",
            "warning": "wymaga wyjaśnienia", 
            "danger": "niespełnione"
        }
        
        assessment_summary = f"""
ANALIZA KAŻDEGO KRYTERIUM:

1. NAGŁOŚĆ ZDARZENIA:
   Ocena: {status_map.get(assessment.sudden.status, assessment.sudden.status)}
   Okoliczności: {assessment.sudden.description}

2. PRZYCZYNA ZEWNĘTRZNA:
   Ocena: {status_map.get(assessment.external.status, assessment.external.status)}
   Okoliczności: {assessment.external.description}

3. ZWIĄZEK Z PRACĄ:
   Ocena: {status_map.get(assessment.work.status, assessment.work.status)}
   Okoliczności: {assessment.work.description}

4. SKUTEK - URAZ LUB ŚMIERĆ:
   Ocena: {status_map.get(assessment.injury.status, assessment.injury.status)}
   Okoliczności: {assessment.injury.description}

TYP DECYZJI: {decision}
- approved = uznanie za wypadek przy pracy
- rejected = odmowa uznania za wypadek przy pracy
- investigation_needed = konieczne postępowanie wyjaśniające
"""

        # Add validation issues if present
        if validation_issues and len(validation_issues) > 0:
            assessment_summary += "\n\nWYKRYTE ROZBIEŻNOŚCI MIĘDZY DOKUMENTAMI:\n"
            for issue in validation_issues:
                severity_pl = {
                    "error": "BŁĄD KRYTYCZNY",
                    "warning": "OSTRZEŻENIE",
                    "info": "INFORMACJA"
                }.get(issue.get("severity", "info"), "INFORMACJA")
                
                assessment_summary += f"\n- {severity_pl} - {issue.get('field', 'Nieznane pole')}: {issue.get('message', 'Brak opisu')}"
                if issue.get("pdfValue") or issue.get("docxValue"):
                    assessment_summary += f"\n  * Wartość w zgłoszeniu ZUS (PDF): {issue.get('pdfValue', 'brak')}"
                    assessment_summary += f"\n  * Wartość w wyjaśnieniu poszkodowanego (DOCX): {issue.get('docxValue', 'brak')}"
            
            assessment_summary += "\n\nUWAGA: Rozbieżności między dokumentami mogą wskazywać na problemy z wiarygodnością zeznań lub błędy w dokumentacji. Należy to uwzględnić w uzasadnieniu, szczególnie przy decyzji o odmowie."
        
        messages = [
            SystemMessage(content=justification_prompt),
            HumanMessage(content=f"Na podstawie poniższej analizy każdego kryterium prawnego oraz typu decyzji, napisz profesjonalne uzasadnienie jako inspektor ds. wypadków przy pracy. Uzasadnienie musi brzmieć naturalnie, jak napisane przez człowieka. Zwróć odpowiedź w formacie JSON z kluczem 'justification'.\n\n{assessment_summary}")
        ]
        
        result = structured_llm.invoke(messages)
        return result.justification
    
    except Exception as e:
        print(f"Error in generate_decision_justification: {e}")
        import traceback
        traceback.print_exc()
        # Return default justification based on decision type
        if decision == "approved":
            return "Po przeprowadzeniu analizy dokumentacji oraz okoliczności zdarzenia, stwierdzam, że przedmiotowe zdarzenie spełnia wszystkie kryteria wypadku przy pracy określone w art. 3 ust. 1 ustawy z dnia 30 października 2002 r. o ubezpieczeniu społecznym z tytułu wypadków przy pracy i chorób zawodowych. Zdarzenie charakteryzowało się nagłością, zostało wywołane przyczyną zewnętrzną i pozostaje w bezpośrednim związku z wykonywaną pracą."
        elif decision == "rejected":
            return "Po analizie przedstawionej dokumentacji oraz okoliczności zdarzenia, stwierdzam, że przedmiotowe zdarzenie nie spełnia kryteriów wypadku przy pracy w rozumieniu art. 3 ust. 1 ustawy z dnia 30 października 2002 r. o ubezpieczeniu społecznym z tytułu wypadków przy pracy i chorób zawodowych. Analiza wykazała brak niezbędnych związków lub niespełnienie innych warunków określonych w przepisach prawa."
        else:
            return "Po wstępnej analizie przedstawionej dokumentacji stwierdzam, że w celu podjęcia ostatecznej decyzji o uznaniu lub odmowie uznania zdarzenia za wypadek przy pracy, konieczne jest przeprowadzenie postępowania wyjaśniającego. Dopiero po otrzymaniu kompletnej dokumentacji możliwe będzie wydanie merytorycznego rozstrzygnięcia w sprawie."


def extract_pdf_data(pdf_file: BinaryIO) -> PDFExtractionResponse:
    """
    Extract structured data from PDF file using LLM with advanced prompt
    """
    try:
        # Read PDF content
        pdf_reader = PdfReader(pdf_file)
        pdf_text = ""
        for page in pdf_reader.pages:
            pdf_text += page.extract_text() + "\n"

        messages = [
            SystemMessage(content=pdf_extraction_prompt),
            HumanMessage(content=f"Wyodrębnij dane z następującego tekstu formularza:\n\n{pdf_text}")
        ]

        try:
            # Use LLM to extract structured data
            structured_llm = llm_llama.with_structured_output(
                schema=PDFExtractionResponse,
                method="json_mode"
            )

            result = structured_llm.invoke(messages)
        except:
            # Fallback to llama model if Groq fails
            structured_llm = llm.with_structured_output(
                schema=PDFExtractionResponse,
                method="json_mode"
            )
            result = structured_llm.invoke(messages)
        return result
    
    except Exception as e:
        print(f"Error in extract_pdf_data: {e}")
        # Return empty response on error
        return PDFExtractionResponse()


def extract_docx_explanation(docx_file: BinaryIO) -> DOCXExplanationResponse:
    """
    Extract structured data from DOCX file with injured person's explanation
    using regex patterns to parse the document structure
    """
    try:
        # Read DOCX content
        doc = Document(docx_file)
        
        # Extract all text from paragraphs and tables
        parts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text and cell.text.strip())
                if row_text:
                    parts.append(row_text)
        
        full_text = "\n".join(parts)
        
        # Initialize response
        response = DOCXExplanationResponse()
        
        # Extract explanation date (dnia dd.mm.rrrr, HH:MM:SS)
        date_match = re.search(r'dnia\s+(\d{1,2})\.(\d{1,2})\.(\d{4}),\s+(\d{1,2}):(\d{2}):(\d{2})', full_text)
        if date_match:
            day, month, year, hour, minute, second = date_match.groups()
            response.explanationDate = f"{year}-{month.zfill(2)}-{day.zfill(2)} {hour.zfill(2)}:{minute}:{second}"
        
        # Extract accident date and time
        accident_match = re.search(r'Zaistniałym\s+w\s+dniu\s+(\d{4})-(\d{2})-(\d{2})\s+ok\.?\s+godz\.?\s+(\d{1,2}):(\d{2})', full_text)
        if accident_match:
            year, month, day, hour, minute = accident_match.groups()
            response.accidentDate = f"{year}-{month}-{day}"
            response.accidentTime = f"{hour.zfill(2)}:{minute}"
        
        # Extract accident location
        location_match = re.search(r'godz\.?\s+\d{1,2}:\d{2}\s+w\s+miejscu\s+(.+?)(?:\n|Poszkodowany)', full_text, re.DOTALL)
        if location_match:
            response.accidentLocation = location_match.group(1).strip()
        
        # Extract personal data
        name_match = re.search(r'Imię\s*\(imiona\)\s*i\s*nazwisko\s+([^\n]+?)\s+Imię\s+ojca', full_text)
        if name_match:
            full_name = name_match.group(1).strip()
            name_parts = full_name.split()
            if len(name_parts) >= 2:
                response.firstName = name_parts[0]
                response.lastName = " ".join(name_parts[1:])
            elif len(name_parts) == 1:
                response.lastName = name_parts[0]
        
        # Father's name
        father_match = re.search(r'Imię\s+ojca\s+([^\n]+)', full_text)
        if father_match:
            response.fatherName = father_match.group(1).strip()
        
        # Birth date and place
        birth_match = re.search(r'Data\s+i\s+miejsce\s+urodzenia\s+(\d{4})-(\d{2})-(\d{2}),\s*([^\n]+)', full_text)
        if birth_match:
            year, month, day, place = birth_match.groups()
            response.birthDate = f"{year}-{month}-{day}"
            response.birthPlace = place.strip()
        
        # PESEL
        pesel_match = re.search(r'Numer\s+PESEL\s+([\d]+)', full_text)
        if pesel_match:
            response.pesel = pesel_match.group(1).strip()
        
        # NIP
        nip_match = re.search(r'NIP\s+([\d]+)', full_text)
        if nip_match:
            response.nip = nip_match.group(1).strip()
        
        # Residence
        residence_match = re.search(r'Miejsce\s+zamieszkania\s+([^\n]+)', full_text)
        if residence_match:
            response.residence = residence_match.group(1).strip()
        
        # Employer information
        employer_match = re.search(r'Miejsce\s+zatrudnienia\s+([^,\n]+)(?:,\s*([^\n]+))?', full_text)
        if employer_match:
            response.employerName = employer_match.group(1).strip()
            if employer_match.group(2):
                response.employerLocation = employer_match.group(2).strip()
        
        # Position
        position_match = re.search(r'Stanowisko\s+lub\s+rodzaj\s+pracy\s+([^\n]+)', full_text)
        if position_match:
            response.position = position_match.group(1).strip()
        
        # Document information
        doc_match = re.search(r'Nazwa\s+i\s+numer\s+dokumentu\s+([^\n]+?)(?:\s+([A-Z\d]+))?,\s*(.+?)(?:\n|wyjaśnia)', full_text, re.DOTALL)
        if doc_match:
            response.documentType = doc_match.group(1).strip()
            if doc_match.group(2):
                response.documentNumber = doc_match.group(2).strip()
            if doc_match.group(3):
                response.documentIssuedBy = doc_match.group(3).strip()
        
        # Explanation text
        explanation_match = re.search(r'wyjaśnia,\s+co\s+następuje:(.+?)(?:Potwierdzeniem|$)', full_text, re.DOTALL)
        if explanation_match:
            response.explanationText = explanation_match.group(1).strip()
        
        # Medical documents
        medical_match = re.search(r'Potwierdzeniem\s+zaistniałych\s+uszkodzeń\s+ciała\s+są\s+następujące\s+dokumenty\s+medyczne:(.+?)(?:\.{5,}|$)', full_text, re.DOTALL)
        if medical_match:
            response.medicalDocuments = medical_match.group(1).strip()
        
        return response
    
    except Exception as e:
        print(f"Error in extract_docx_explanation: {e}")
        # Return empty response on error
        return DOCXExplanationResponse()


def format_date(date_str: str) -> str:
    """Format date from YYYY-MM-DD to DDMMYYYY"""
    if not date_str or len(date_str) != 10:
        return date_str
    try:
        year, month, day = date_str.split('-')
        return f"{day}{month}{year}"
    except:
        return date_str


def fill_pdf_with_pymupdf(input_pdf_path: str, json_data: dict) -> bytes:
    """Fill PDF form using PyMuPDF and return bytes"""
    doc = fitz.open(input_pdf_path)

    # Page1 - dane osobowe + ostatni adres w Polsce
    field_data = {
        'PESEL[0]': json_data.get('pesel', ''),
        'Imię[0]': json_data.get('firstName', ''),
        'Nazwisko[0]': json_data.get('lastName', ''),
        'Dataurodzenia[0]': format_date(json_data.get('birthDate', '')),
        'Miejsceurodzenia[0]': json_data.get('birthPlace', ''),
        'Numertelefonu[0]': json_data.get('phoneNumber', ''),
        'Rodzajseriainumerdokumentu[0]': f"{json_data.get('documentType', '')} {json_data.get('documentSeries', '')} {json_data.get('documentNumber', '')}",
        'Ulica[0]': json_data.get('street', ''),
        'Numerdomu[0]': json_data.get('houseNumber', ''),
        'Numerlokalu[0]': json_data.get('apartmentNumber', ''),
        'Kodpocztowy[0]': json_data.get('postalCode', ''),
        'Poczta[0]': json_data.get('city', ''),
        'Nazwapaństwa[0]': json_data.get('country', ''),
        # Ostatni adres w Polsce
        'Ulica2A[0]': json_data.get('lastPolandStreet', ''),
        'Numerdomu2A[0]': json_data.get('lastPolandHouseNumber', ''),
        'Numerlokalu2A[0]': json_data.get('lastPolandApartmentNumber', ''),
        'Kodpocztowy2A[0]': json_data.get('lastPolandPostalCode', ''),
        'Poczta2A[0]': json_data.get('lastPolandCity', ''),
        # Pozostałe pola
        'Datawyp[0]': format_date(json_data.get('accidentDate', '')),
        'Godzina[0]': json_data.get('accidentTime', ''),
        'Miejscewyp[0]': json_data.get('accidentLocation', ''),
        'Godzina3A[0]': json_data.get('plannedStartTime', ''),
        'Godzina3B[0]': json_data.get('plannedEndTime', ''),
        'Tekst4[0]': json_data.get('injuryType', ''),
        'Tekst5[0]': json_data.get('accidentDescription', ''),
        'Tekst6[0]': json_data.get('healthFacilityInfo', ''),
        'Tekst7[0]': json_data.get('investigatingAuthority', ''),
        'Tekst8[0]': json_data.get('machineryCondition', ''),
        'Data[0]': format_date(json_data.get('declarationDate', '')),
        'Data[1]': format_date(json_data.get('documentsDeliveryDate', '')),
        'Inne[0]': json_data.get('otherAttachments', ''),
    }

    # Page2 - adres korespondencyjny
    field_last_corr = {
        'Ulica2[1]': json_data.get('corrStreet', ''),
        'Numerdomu2[1]': json_data.get('corrHouseNumber', ''),
        'Numerlokalu2[1]': json_data.get('corrApartmentNumber', ''),
        'Kodpocztowy2[1]': json_data.get('corrPostalCode', ''),
        'Poczta2[1]': json_data.get('corrCity', ''),
        'Nazwapaństwa2[0]': json_data.get('corrCountry', ''),
    }

    # Checkboxes
    response_method = json_data.get('responseMethod', '').lower()
    check_placowka = 'w placówce zus' in response_method
    check_poczta = 'pocztą na adres' in response_method
    check_pue = 'pue zus' in response_method

    checkbox_data = {
        'TAK6[0]': json_data.get('wasFirstAidGiven') == 'Tak',
        'NIE6[0]': json_data.get('wasFirstAidGiven') == 'Nie',
        'TAK8[0]': json_data.get('wasMachineryInvolved') == 'Tak',
        'NIE8[0]': json_data.get('wasMachineryInvolved') == 'Nie',
        'TAK9[0]': json_data.get('hasCertification') == 'Tak',
        'NIE9[0]': json_data.get('hasCertification') == 'Nie',
        'TAK10[0]': json_data.get('isInInventory') == 'Tak',
        'NIE10[0]': json_data.get('isInInventory') == 'Nie',
        'ZaznaczX1[0]': json_data.get('attachHospitalCard') == True,
        'ZaznaczX2[0]': json_data.get('attachProsecutorDecision') == True,
        'ZaznaczX3[0]': json_data.get('attachDeathCertificate') == True,
        'ZaznaczX4[0]': json_data.get('attachRightToIssueCard') == True,
        'wplacowce[0]': check_placowka,
        'poczta[0]': check_poczta,
        'PUE[0]': check_pue,
        'adres[0]': json_data.get('isLastPolandCorrespondenceAddress', '').lower() == 'tak',
    }

    for page in doc:
        for widget in page.widgets():
            field_name = widget.field_name
            field_type = widget.field_type

            # Text fields
            for key, value in {**field_data, **field_last_corr}.items():
                if key in field_name and value and str(value).strip():
                    try:
                        widget.field_value = str(value)
                        widget.text_encoding = fitz.TEXT_ENCODING_LATIN
                        widget.update()
                    except:
                        pass

            # Checkboxes
            for key, should_check in checkbox_data.items():
                if key in field_name and field_type == 2:
                    try:
                        if should_check:
                            rect = widget.rect
                            margin = 3
                            page.draw_line((rect.x0 + margin, rect.y0 + margin),
                                           (rect.x1 - margin, rect.y1 - margin),
                                           color=(0, 0, 0), width=2.5)
                            page.draw_line((rect.x1 - margin, rect.y0 + margin),
                                           (rect.x0 + margin, rect.y1 - margin),
                                           color=(0, 0, 0), width=2.5)
                        else:
                            widget.field_value = False
                            widget.update()
                    except:
                        pass

    # Save to bytes
    pdf_bytes = doc.tobytes(garbage=4, deflate=True, clean=True, pretty=True)
    doc.close()
    return pdf_bytes


def generate_accident_notification_pdf(data: AccidentNotificationRequest) -> bytes:
    """
    Generate filled accident notification PDF from form data
    Returns PDF as bytes
    """
    input_pdf = 'EWYP_wypelnij_i_wydrukuj.pdf'
    
    if not os.path.exists(input_pdf):
        raise FileNotFoundError(f"Template PDF not found: {input_pdf}")
    
    # Convert Pydantic model to dict
    json_data = data.model_dump()
    
    return fill_pdf_with_pymupdf(input_pdf, json_data)


def replace_placeholders_in_paragraph(paragraph, mapping):
    """Replace placeholders {{key}} in a single paragraph"""
    # Get full paragraph text
    full_text = paragraph.text
    
    # Replace all placeholders in the full text
    for key, val in mapping.items():
        placeholder = f"{{{{{key}}}}}"
        full_text = full_text.replace(placeholder, str(val))
    
    # If no replacements were made, return early
    if full_text == paragraph.text:
        return
    
    # Clear existing runs and add new text
    # First, preserve the formatting from the first run
    if paragraph.runs:
        first_run = paragraph.runs[0]
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        # Set the new text in the first run
        first_run.text = full_text
    else:
        # If no runs exist, add the text
        paragraph.add_run(full_text)


def replace_placeholders_in_tables(doc, mapping):
    """Replace placeholders in tables"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, mapping)


def generate_injured_statement_docx(data: InjuredStatementRequest) -> bytes:
    """
    Generate filled injured person's statement DOCX from form data
    Returns DOCX as bytes
    """
    template_path = 'B1a_wyjasnienia_poszkodowanego_o_wypadku_przy_pracy.docx'
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template DOCX not found: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    # Convert Pydantic model to dict
    json_data = data.model_dump()
    
    # Prepare mapping for placeholders
    mapping = {
        "generatedDate": json_data.get("generatedDate", ""),
        "accidentDate": json_data.get("accidentDate", ""),
        "accidentTime": json_data.get("accidentTime", ""),
        "accidentLocation":
            f"{json_data.get('accidentStreet', '')} {json_data.get('accidentHouseNumber', '')}/"
            f"{json_data.get('accidentApartmentNumber', '')}, "
            f"{json_data.get('accidentPostalCode', '')} {json_data.get('accidentCity', '')}",
        
        "firstNameLastName": f"{json_data.get('firstName', '')} {json_data.get('lastName', '')}",
        "fatherName": json_data.get("fatherName", ""),
        "birthDatePlace": f"{json_data.get('birthDate', '')}, {json_data.get('birthPlace', '')}",
        "pesel": json_data.get("pesel", ""),
        "nip": json_data.get("nip", ""),
        "residenceAddress": json_data.get("residenceAddress", ""),
        "employmentPlace": json_data.get("employmentPlace", ""),
        "position": json_data.get("position", ""),
        "identityDocument": json_data.get("identityDocument", ""),
        "accidentDescription": json_data.get("accidentDescription", ""),
        "medicalDocuments": ", ".join(json_data.get("medicalDocuments", [])) if json_data.get("medicalDocuments") else ""
    }
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, mapping)
    
    # Replace placeholders in tables
    replace_placeholders_in_tables(doc, mapping)


def compare_pdf_docx_data(pdf_data: PDFExtractionResponse, docx_data: DOCXExplanationResponse) -> DataComparisonResponse:
    """
    Compare data extracted from PDF (ZUS form) and DOCX (injured person's explanation)
    Validate consistency and merge data with priority
    """
    issues: List[ValidationIssue] = []
    
    # Helper function to normalize strings for comparison
    def normalize(s: str) -> str:
        return s.strip().lower() if s else ""
    
    # 1. Compare PESEL
    if pdf_data.pesel and docx_data.pesel:
        if normalize(pdf_data.pesel) != normalize(docx_data.pesel):
            issues.append(ValidationIssue(
                field="pesel",
                severity="error",
                message="Numer PESEL różni się między dokumentami",
                pdfValue=pdf_data.pesel,
                docxValue=docx_data.pesel
            ))
    
    # 2. Compare first name
    if pdf_data.firstName and docx_data.firstName:
        if normalize(pdf_data.firstName) != normalize(docx_data.firstName):
            issues.append(ValidationIssue(
                field="firstName",
                severity="warning",
                message="Imię różni się między dokumentami",
                pdfValue=pdf_data.firstName,
                docxValue=docx_data.firstName
            ))
    
    # 3. Compare last name
    if pdf_data.lastName and docx_data.lastName:
        if normalize(pdf_data.lastName) != normalize(docx_data.lastName):
            issues.append(ValidationIssue(
                field="lastName",
                severity="warning",
                message="Nazwisko różni się między dokumentami",
                pdfValue=pdf_data.lastName,
                docxValue=docx_data.lastName
            ))
    
    # 4. Compare birth date
    if pdf_data.birthDate and docx_data.birthDate:
        if pdf_data.birthDate != docx_data.birthDate:
            issues.append(ValidationIssue(
                field="birthDate",
                severity="error",
                message="Data urodzenia różni się między dokumentami",
                pdfValue=pdf_data.birthDate,
                docxValue=docx_data.birthDate
            ))
    
    # 5. Compare birth place
    if pdf_data.birthPlace and docx_data.birthPlace:
        if normalize(pdf_data.birthPlace) != normalize(docx_data.birthPlace):
            issues.append(ValidationIssue(
                field="birthPlace",
                severity="info",
                message="Miejsce urodzenia różni się między dokumentami",
                pdfValue=pdf_data.birthPlace,
                docxValue=docx_data.birthPlace
            ))
    
    # 6. Compare document type
    if pdf_data.documentType and docx_data.documentType:
        pdf_doc_normalized = normalize(pdf_data.documentType)
        docx_doc_normalized = normalize(docx_data.documentType)
        if pdf_doc_normalized not in docx_doc_normalized and docx_doc_normalized not in pdf_doc_normalized:
            issues.append(ValidationIssue(
                field="documentType",
                severity="warning",
                message="Rodzaj dokumentu tożsamości różni się między dokumentami",
                pdfValue=pdf_data.documentType,
                docxValue=docx_data.documentType
            ))
    
    # 7. Compare document number
    if pdf_data.documentNumber and docx_data.documentNumber:
        pdf_num = normalize(pdf_data.documentNumber.replace(" ", ""))
        docx_num = normalize(docx_data.documentNumber.replace(" ", ""))
        if pdf_num != docx_num:
            issues.append(ValidationIssue(
                field="documentNumber",
                severity="warning",
                message="Numer dokumentu różni się między dokumentami",
                pdfValue=pdf_data.documentNumber,
                docxValue=docx_data.documentNumber
            ))
    
    # 8. CRITICAL: Compare accident date
    if pdf_data.accidentDate and docx_data.accidentDate:
        if pdf_data.accidentDate != docx_data.accidentDate:
            issues.append(ValidationIssue(
                field="accidentDate",
                severity="error",
                message="KRYTYCZNE: Data wypadku różni się między dokumentami!",
                pdfValue=pdf_data.accidentDate,
                docxValue=docx_data.accidentDate
            ))
    elif not pdf_data.accidentDate or not docx_data.accidentDate:
        issues.append(ValidationIssue(
            field="accidentDate",
            severity="error",
            message="Data wypadku nie została podana w jednym z dokumentów",
            pdfValue=pdf_data.accidentDate or "",
            docxValue=docx_data.accidentDate or ""
        ))
    
    # 9. Compare accident time
    if pdf_data.accidentTime and docx_data.accidentTime:
        if pdf_data.accidentTime != docx_data.accidentTime:
            issues.append(ValidationIssue(
                field="accidentTime",
                severity="warning",
                message="Godzina wypadku różni się między dokumentami",
                pdfValue=pdf_data.accidentTime,
                docxValue=docx_data.accidentTime
            ))
    
    # 10. Compare accident location
    if pdf_data.accidentLocation and docx_data.accidentLocation:
        if normalize(pdf_data.accidentLocation) != normalize(docx_data.accidentLocation):
            issues.append(ValidationIssue(
                field="accidentLocation",
                severity="info",
                message="Miejsce wypadku różni się między dokumentami (może być bardziej szczegółowe w wyjaśnieniach)",
                pdfValue=pdf_data.accidentLocation,
                docxValue=docx_data.accidentLocation
            ))
    
    # Merge data with priority (PDF has priority for official data, DOCX adds explanation details)
    merged_data = {
        # Personal data - priority: PDF (official form)
        "pesel": pdf_data.pesel or docx_data.pesel,
        "firstName": pdf_data.firstName or docx_data.firstName,
        "lastName": pdf_data.lastName or docx_data.lastName,
        "birthDate": pdf_data.birthDate or docx_data.birthDate,
        "birthPlace": pdf_data.birthPlace or docx_data.birthPlace,
        "phoneNumber": pdf_data.phoneNumber,
        
        # Document data - priority: PDF
        "documentType": pdf_data.documentType or docx_data.documentType,
        "documentSeries": pdf_data.documentSeries,
        "documentNumber": pdf_data.documentNumber or docx_data.documentNumber,
        
        # Address - from PDF
        "street": pdf_data.street,
        "houseNumber": pdf_data.houseNumber,
        "apartmentNumber": pdf_data.apartmentNumber,
        "postalCode": pdf_data.postalCode,
        "city": pdf_data.city,
        "country": pdf_data.country,
        "isCorrespondenceAddress": pdf_data.isCorrespondenceAddress,
        
        # Business activity
        "hasBusinessActivity": pdf_data.hasBusinessActivity,
        
        # Notifier data - from PDF
        "notifierFirstName": pdf_data.notifierFirstName,
        "notifierLastName": pdf_data.notifierLastName,
        
        # Accident data - priority: PDF for official data, DOCX for descriptions
        "accidentDate": pdf_data.accidentDate or docx_data.accidentDate,
        "accidentTime": pdf_data.accidentTime or docx_data.accidentTime,
        "accidentLocation": pdf_data.accidentLocation or docx_data.accidentLocation,
        "plannedStartTime": pdf_data.plannedStartTime,
        "plannedEndTime": pdf_data.plannedEndTime,
        "injuryType": pdf_data.injuryType,
        
        # Description - merge both sources
        "accidentDescription": pdf_data.accidentDescription or docx_data.explanationText,
        "detailedExplanation": docx_data.explanationText,  # Additional field from DOCX
        
        # Medical and investigation data - from PDF
        "wasFirstAidGiven": pdf_data.wasFirstAidGiven,
        "healthFacilityInfo": pdf_data.healthFacilityInfo,
        "investigatingAuthority": pdf_data.investigatingAuthority,
        "wasMachineryInvolved": pdf_data.wasMachineryInvolved,
        "machineryCondition": pdf_data.machineryCondition,
        "hasCertification": pdf_data.hasCertification,
        "isInInventory": pdf_data.isInInventory,
        
        # Witnesses - from PDF
        "witness1": pdf_data.witness1.model_dump() if pdf_data.witness1 else {},
        "witness2": pdf_data.witness2.model_dump() if pdf_data.witness2 else {},
        "witness3": pdf_data.witness3.model_dump() if pdf_data.witness3 else {},
        
        # Attachments - from PDF
        "attachHospitalCard": pdf_data.attachHospitalCard,
        "attachProsecutorDecision": pdf_data.attachProsecutorDecision,
        "attachDeathCertificate": pdf_data.attachDeathCertificate,
        "attachRightToIssueCard": pdf_data.attachRightToIssueCard,
        "otherAttachments": pdf_data.otherAttachments,
        "documentsDeliveryDate": pdf_data.documentsDeliveryDate,
        "additionalDocuments": pdf_data.additionalDocuments,
        
        # Additional data from DOCX
        "employerName": docx_data.employerName,
        "employerLocation": docx_data.employerLocation,
        "position": docx_data.position,
        "fatherName": docx_data.fatherName,
        "nip": docx_data.nip,
        "residence": docx_data.residence,
        "documentIssuedBy": docx_data.documentIssuedBy,
        "medicalDocuments": docx_data.medicalDocuments,
        "explanationDate": docx_data.explanationDate,
    }
    
    # Determine if data is valid
    error_count = sum(1 for issue in issues if issue.severity == "error")
    is_valid = error_count == 0
    
    # Generate summary
    if is_valid:
        summary = f"Dane są spójne. Znaleziono {len(issues)} uwag(i) do przejrzenia."
    else:
        summary = f"UWAGA: Znaleziono {error_count} krytycznych błędów! Dane wymagają weryfikacji przed przetworzeniem."
    
    return DataComparisonResponse(
        isValid=is_valid,
        issues=issues,
        mergedData=merged_data,
        summary=summary
    )


def add_run_with_strike(paragraph, text, strike=False):
    """Add text run with optional strikethrough"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    run = paragraph.add_run(text)
    if strike:
        r = run._r
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        strike_tag = OxmlElement('w:strike')
        rPr.append(strike_tag)


def generate_accident_card_docx(data: AccidentCardRequest) -> bytes:
    """
    Generate filled accident card DOCX from form data
    Returns DOCX as bytes
    """
    template_path = 'karta-wypadku.docx'
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template DOCX not found: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    # Convert Pydantic model to dict
    json_data = data.model_dump()
    
    # Build victim name
    json_data["victimName"] = f"{json_data.get('victimFirstName', '')} {json_data.get('victimLastName', '')}".strip()
    
    # Handle accidentClause based on accidentIsWorkAccident
    accident_state = json_data.get("accidentIsWorkAccident", "").lower()
    
    if accident_state == "tak":
        accident_clause = {
            "jest": {"text": "Wypadek jest", "strike": False},
            "nie_jest": {"text": "nie jest", "strike": True}
        }
    elif accident_state == "nie":
        accident_clause = {
            "jest": {"text": "Wypadek jest", "strike": True},
            "nie_jest": {"text": "nie jest", "strike": False}
        }
    else:
        accident_clause = {
            "jest": {"text": "Wypadek jest", "strike": False},
            "nie_jest": {"text": "nie jest", "strike": False}
        }
    
    # Handle victimViolationClause
    state = json_data.get("victimViolationProved", "").lower()
    
    if state == "przyznano":
        victim_violation_clause = {
            "tak": {"text": "TAK", "strike": False},
            "nie": {"text": "NIE", "strike": True},
            "suffix": ""
        }
    elif state == "odrzucono":
        victim_violation_clause = {
            "tak": {"text": "TAK", "strike": True},
            "nie": {"text": "NIE", "strike": False},
            "suffix": " (odrzucono)"
        }
    else:  # "nie"
        victim_violation_clause = {
            "tak": {"text": "TAK", "strike": True},
            "nie": {"text": "NIE", "strike": False},
            "suffix": ""
        }
    
    # Replace plain placeholders in paragraphs
    for p in doc.paragraphs:
        for key, val in json_data.items():
            if key not in ["accidentClause", "victimViolationClause"]:
                p.text = p.text.replace("{{" + key + "}}", str(val))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in json_data.items():
                        if key not in ["accidentClause", "victimViolationClause"]:
                            p.text = p.text.replace("{{" + key + "}}", str(val))
    
    # Handle accidentClause with strikethrough
    for p in doc.paragraphs:
        if "{{accidentClause}}" in p.text:
            p.text = ""
            add_run_with_strike(p, accident_clause["jest"]["text"], accident_clause["jest"]["strike"])
            p.add_run(" / ")
            add_run_with_strike(p, accident_clause["nie_jest"]["text"], accident_clause["nie_jest"]["strike"])
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{accidentClause}}" in p.text:
                        p.text = ""
                        add_run_with_strike(p, accident_clause["jest"]["text"], accident_clause["jest"]["strike"])
                        p.add_run(" / ")
                        add_run_with_strike(p, accident_clause["nie_jest"]["text"], accident_clause["nie_jest"]["strike"])
    
    # Handle victimViolationClause with strikethrough
    for p in doc.paragraphs:
        if "{{victimViolationClause}}" in p.text:
            p.text = ""
            add_run_with_strike(p, victim_violation_clause["tak"]["text"], victim_violation_clause["tak"]["strike"])
            p.add_run(" / ")
            add_run_with_strike(p, victim_violation_clause["nie"]["text"], victim_violation_clause["nie"]["strike"])
            p.add_run(victim_violation_clause["suffix"])
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{victimViolationClause}}" in p.text:
                        p.text = ""
                        add_run_with_strike(p, victim_violation_clause["tak"]["text"], victim_violation_clause["tak"]["strike"])
                        p.add_run(" / ")
                        add_run_with_strike(p, victim_violation_clause["nie"]["text"], victim_violation_clause["nie"]["strike"])
                        p.add_run(victim_violation_clause["suffix"])
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()